"""
Word document (.docx) export for approved tender synopses.
Uses python-docx with SAP Blue branding.
"""

from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SAP_BLUE  = RGBColor(0x00, 0x70, 0xC0)
DARK_GREY = RGBColor(0x40, 0x40, 0x40)


def _section_heading(doc, text: str):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = SAP_BLUE
    return p


def _field_block(doc, label: str, value: str):
    lp = doc.add_paragraph()
    r = lp.add_run(label + ":")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = SAP_BLUE
    vp = doc.add_paragraph(value or "Not found in document")
    vp.paragraph_format.left_indent = Pt(14)
    for run in vp.runs:
        run.font.size = Pt(11)
    doc.add_paragraph()


def save_synopsis_docx(synopsis: dict, sourcing_project_id: str, output_path: str):
    """Generate and save a formatted .docx for the approved synopsis."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover ──────────────────────────────────────────────────────────────
    title = doc.add_heading("Tender Synopsis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = SAP_BLUE

    portal = synopsis.get("portalName", "Generic")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Sourcing Project: {sourcing_project_id}  |  "
        f"Portal: {portal}  |  "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"Language: {synopsis.get('language', 'English')}"
    )
    mr.font.size = Pt(10)
    mr.italic = True
    mr.font.color.rgb = DARK_GREY
    doc.add_paragraph()

    # ── Title + Summary ────────────────────────────────────────────────────
    _section_heading(doc, "1. Tender Title")
    _field_block(doc, "Title", synopsis.get("tenderTitle"))

    _section_heading(doc, "2. Executive Summary")
    _field_block(doc, "Summary", synopsis.get("executiveSummary"))

    # ── Dynamic supplierFields grouped by category ─────────────────────────
    supplier_fields = synopsis.get("supplierFields", [])
    category_titles = {
        "overview":    "3. Tender Overview",
        "commercial":  "4. Commercial Details",
        "dates":       "5. Key Dates",
        "eligibility": "6. Eligibility & Qualification",
    }
    for cat, heading_text in category_titles.items():
        cat_fields = [f for f in supplier_fields if f.get("category") == cat]
        if not cat_fields:
            continue
        _section_heading(doc, heading_text)
        for f in cat_fields:
            label   = f.get("label", "Field")
            value   = f.get("value", "Not specified")
            sap_src = f.get("sapSource", "")
            lp = doc.add_paragraph()
            lr = lp.add_run(label + ":")
            lr.bold = True
            lr.font.size = Pt(11)
            lr.font.color.rgb = SAP_BLUE
            vp = doc.add_paragraph(value)
            vp.paragraph_format.left_indent = Pt(14)
            for run in vp.runs:
                run.font.size = Pt(11)
                if value.lower().startswith("not spec"):
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            if sap_src:
                sp = doc.add_paragraph(f"  [SAP: {sap_src}]")
                sp.paragraph_format.left_indent = Pt(14)
                for run in sp.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        doc.add_paragraph()

    # ── Supplier Actions ───────────────────────────────────────────────────
    actions = [a for a in synopsis.get("supplierActions", []) if a]
    if actions:
        _section_heading(doc, "7. Supplier Actions")
        for action in actions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(action).font.size = Pt(11)
        doc.add_paragraph()

    # ── Portal compliance note ─────────────────────────────────────────────
    pcn = synopsis.get("portalComplianceNote", "")
    if pcn:
        _section_heading(doc, "8. Portal Compliance Note")
        _field_block(doc, portal, pcn)

    # ── Portal missing fields ──────────────────────────────────────────────
    pmf = synopsis.get("portalMissingFields", [])
    if pmf:
        _section_heading(doc, "Required Portal Fields Not in SAP")
        items = pmf if (pmf and isinstance(pmf[0], dict)) else [{"label": m, "reason": ""} for m in pmf]
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            label  = item.get("label", str(item))
            reason = item.get("reason", "")
            r = p.add_run(f"{label}" + (f" — {reason}" if reason else ""))
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            r.font.size = Pt(11)
        doc.add_paragraph()

    # ── Missing SAP info ───────────────────────────────────────────────────
    missing = [m for m in synopsis.get("missingInformation", []) if m]
    if missing:
        _section_heading(doc, "Flags — Missing SAP Information")
        for item in missing:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(item)
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            r.font.size = Pt(11)
        doc.add_paragraph()

    # ── Source references ──────────────────────────────────────────────────
    refs = [r for r in synopsis.get("sourceReferences", []) if r]
    if refs:
        _section_heading(doc, "SAP PPS Source Fields Used")
        for ref in refs:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(ref).font.size = Pt(10)
        doc.add_paragraph()

    # ── Footer ─────────────────────────────────────────────────────────────
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        "AI-generated from SAP PPS EMT 601 (UI_SOURCINGPROJECT_MANAGE_2). "
        "Approved by procurement officer before publication."
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = DARK_GREY

    doc.save(output_path)
